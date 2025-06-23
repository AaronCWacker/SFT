import gradio as gr
from pathlib import Path
import datetime
import re
import os
import shutil
import io
import base64
from collections import defaultdict
from PIL import Image

# Document Generation Libs
from docx import Document
import openpyxl
from pypdf import PdfWriter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, BaseDocTemplate, Frame, PageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# AI and Media Libs
from openai import AzureOpenAI
import fitz  # PyMuPDF

# --- Configuration & Setup ---
CWD = Path.cwd()
OUTPUT_DIR = CWD / "generated_outputs"
PREVIEW_DIR = CWD / "previews"
FONT_DIR = CWD
OUTPUT_DIR.mkdir(exist_ok=True)
PREVIEW_DIR.mkdir(exist_ok=True)

LAYOUTS = {
    "A4 Portrait": {"size": A4},
    "A4 Landscape": {"size": landscape(A4)},
    "Letter Portrait": {"size": letter},
    "Letter Landscape": {"size": landscape(letter)},
}

# 🧠 Initialize Azure OpenAI Client
# NOTE: This requires AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY in your environment.
try:
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_version="2024-05-01-preview",
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    )
    AZURE_CLIENT_AVAILABLE = True
except Exception as e:
    print("Warning: Azure OpenAI client could not be initialized. Text generation will use dummy data.")
    print(f"Error: {e}")
    client = None
    AZURE_CLIENT_AVAILABLE = False

# 📖 Map UI model names to your actual Azure deployment names.
# YOU MUST CHANGE THESE DEPLOYMENT NAMES to match your Azure setup.
AZURE_DEPLOYMENT_NAMES = {
    # Chat / Vision Models
    "gpt-4o": "your-gpt-4o-deployment-name",
    "gpt-4.1": "your-gpt-4.1-deployment-name",
    "gpt-4.1-mini": "your-gpt-4.1-mini-deployment-name",
    "gpt-4o-mini": "your-gpt-4o-mini-deployment-name",
    "gpt-4o-realtime-preview": "your-gpt-4o-realtime-deployment-name",
    # Reasoning Models
    "o1-mini": "your-o1-mini-deployment-name",
    "o3-mini": "your-o3-mini-deployment-name",
    "o4-mini": "your-o4-mini-deployment-name",
    # Transcription Models
    "gpt-4o-transcribe": "your-gpt-4o-transcribe-deployment",
    "gpt-4o-mini-transcribe": "your-gpt-4o-mini-transcribe-deployment",
}


# --- ✍️ Document Generation Engines ---

def create_pdf(md_content, font_name, emoji_font, pagesize, num_columns):
    """📄 Builds a beautiful PDF from a Markdown story using ReportLab."""
    pdf_buffer = io.BytesIO()
    story = markdown_to_story(md_content, font_name, emoji_font)
    if num_columns > 1:
        doc = BaseDocTemplate(pdf_buffer, pagesize=pagesize, leftMargin=0.5 * inch, rightMargin=0.5 * inch)
        frame_width = (doc.width / num_columns) - (num_columns - 1) * 0.1 * inch
        frames = [Frame(doc.leftMargin + i * (frame_width + 0.2 * inch), doc.bottomMargin, frame_width, doc.height) for i in range(num_columns)]
        doc.addPageTemplates([PageTemplate(id='MultiCol', frames=frames)])
    else:
        doc = SimpleDocTemplate(pdf_buffer, pagesize=pagesize)
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

def create_docx(md_content):
    """📝 Crafts a DOCX document, translating Markdown to Word elements."""
    document = Document()
    for line in md_content.split('\n'):
        if line.startswith('# '): document.add_heading(line[2:], level=1)
        elif line.startswith('## '): document.add_heading(line[3:], level=2)
        elif line.strip().startswith(('- ', '* ')): document.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            p = document.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'): p.add_run(part[2:-2]).bold = True
                else: p.add_run(part)
    return document

def create_xlsx(md_content):
    """📊 Organizes a Markdown outline into columns in an XLSX file."""
    workbook = openpyxl.Workbook(); sheet = workbook.active
    sections = re.split(r'\n# ', '\n' + md_content)
    if sections and sections[0] == '': sections.pop(0)
    column_data = []
    for section in sections:
        lines = section.split('\n'); header = lines[0]
        content = [l.strip() for l in lines[1:] if l.strip()]
        column_data.append({'header': header, 'content': content})
    for c_idx, col in enumerate(column_data, 1):
        sheet.cell(row=1, column=c_idx, value=col['header'])
        for r_idx, line_content in enumerate(col['content'], 2):
            sheet.cell(row=r_idx, column=c_idx, value=line_content)
    return workbook

def markdown_to_story(markdown_text: str, font_name: str, emoji_font: str):
    """📜 Translates Markdown text into a sequence of ReportLab flowables for PDF rendering."""
    styles = getSampleStyleSheet()
    bold_font = f"{font_name}-Bold" if font_name != "Helvetica" else "Helvetica-Bold"
    style_normal = ParagraphStyle('BodyText', fontName=font_name, spaceAfter=6, fontSize=10)
    style_h1 = ParagraphStyle('h1', fontName=bold_font, spaceBefore=12, fontSize=24)
    story, first_heading = [], True
    for line in markdown_text.split('\n'):
        content, style = line, style_normal
        if line.startswith("# "):
            if not first_heading: story.append(PageBreak())
            content, style, first_heading = line.lstrip('# '), style_h1, False
        formatted_content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
        final_content = apply_emoji_font(formatted_content, emoji_font)
        story.append(Paragraph(final_content, style))
    return story


# --- 🔮 Omni-Model Processing ---

def process_text_input(prompt, model_deployment_name):
    """💬 Sends a text prompt to the Azure OpenAI model and gets a response."""
    if not AZURE_CLIENT_AVAILABLE: return "Azure OpenAI client not configured. This is dummy text."
    completion = client.chat.completions.create(
        model=model_deployment_name,
        messages=[{"role": "user", "content": prompt}]
    )
    return completion.choices[0].message.content

def process_image_input(image_file, prompt, model_deployment_name):
    """🖼️ Encodes an image and sends it with a prompt to the Azure OpenAI model."""
    if not AZURE_CLIENT_AVAILABLE: return "Azure OpenAI client not configured. This is a dummy image description."
    with Image.open(image_file.name) as img:
        with io.BytesIO() as output:
            img.save(output, format="PNG")
            base64_image = base64.b64encode(output.getvalue()).decode("utf-8")
    
    response = client.chat.completions.create(
        model=model_deployment_name,
        messages=[{"role": "user", "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}}
        ]}]
    )
    return response.choices[0].message.content

def process_audio_input(audio_file, prompt, chat_model_deployment, transcribe_model_deployment):
    """🎤 Transcribes audio and sends the text with a prompt to the Azure OpenAI model."""
    if not AZURE_CLIENT_AVAILABLE: return "Azure OpenAI client not configured. This is a dummy audio summary."
    with open(audio_file.name, "rb") as f:
        transcription = client.audio.transcriptions.create(
            model=transcribe_model_deployment,
            file=f
        ).text
    
    full_prompt = f"{prompt}\n\nAudio Transcription:\n{transcription}"
    return process_text_input(full_prompt, chat_model_deployment)

def process_pdf_input(pdf_file, prompt, model_deployment_name, progress):
    """📄 Performs OCR on a PDF by sending pages as images to the AI model."""
    if not AZURE_CLIENT_AVAILABLE: return "Azure OpenAI client not configured. This is a dummy PDF summary."
    
    all_extracted_text = []
    doc = fitz.open(pdf_file.name)
    
    # Process pages in pairs
    for i in progress.tqdm(range(0, len(doc), 2), desc="Performing PDF OCR"):
        page_images = []
        messages = [{"type": "text", "text": prompt}]

        # Get first page of the pair
        page1 = doc.load_page(i)
        pix1 = page1.get_pixmap(dpi=150)
        img_bytes1 = pix1.tobytes("png")
        base64_image1 = base64.b64encode(img_bytes1).decode("utf-8")
        messages.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image1}"}})

        # Get second page if it exists
        if i + 1 < len(doc):
            page2 = doc.load_page(i + 1)
            pix2 = page2.get_pixmap(dpi=150)
            img_bytes2 = pix2.tobytes("png")
            base64_image2 = base64.b64encode(img_bytes2).decode("utf-8")
            messages.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image2}"}})
        
        response = client.chat.completions.create(
            model=model_deployment_name,
            messages=[{"role": "user", "content": messages}]
        )
        all_extracted_text.append(response.choices[0].message.content)

    return "\n\n".join(all_extracted_text)


# --- 🛠️ Helpers & Main API ---

def register_local_fonts():
    """✒️ Scans for local .ttf fonts and registers them for PDF creation."""
    text_font_names, emoji_font_name = [], None
    font_files = list(FONT_DIR.glob("*.ttf"))
    for font_path in font_files:
        try:
            font_name = font_path.stem
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            pdfmetrics.registerFont(TTFont(f"{font_name}-Bold", str(font_path)))
            pdfmetrics.registerFontFamily(font_name, normal=font_name, bold=f"{font_name}-Bold")
            if "notocoloremoji-regular" in font_name.lower():
                emoji_font_name = font_name
            else:
                text_font_names.append(font_name)
        except Exception as e:
            print(f"Could not register font {font_path.name}: {e}")
    if not text_font_names: text_font_names.append('Helvetica')
    return sorted(text_font_names), emoji_font_name

def apply_emoji_font(text: str, emoji_font_name: str) -> str:
    """😊 Finds emojis and wraps them in special font tags for the PDF."""
    if not emoji_font_name: return text
    emoji_pattern = re.compile(f"([{re.escape(''.join(map(chr, range(0x1f600, 0x1f650))))}"
                               f"{re.escape(''.join(map(chr, range(0x1f300, 0x1f5ff))))}]+)")
    return emoji_pattern.sub(fr'<font name="{emoji_font_name}">\1</font>', text)

def create_pdf_preview(pdf_path: Path):
    """🏞️ Generates a PNG thumbnail for the first page of a PDF."""
    preview_path = PREVIEW_DIR / f"{pdf_path.stem}.png"
    try:
        doc = fitz.open(pdf_path); page = doc.load_page(0); pix = page.get_pixmap()
        pix.save(str(preview_path)); doc.close()
        return str(preview_path)
    except: return None

def generate_outputs_api(omni_files, omni_prompt, chat_model, transcribe_model, output_formats, layouts, fonts, num_columns, page_w_mult, page_h_mult, progress=gr.Progress(track_tqdm=True)):
    """🚀 The main entry point that orchestrates the entire multi-modal generation process."""
    if not omni_prompt and not omni_files: raise gr.Error("Please provide a prompt or upload at least one file.")
    if not output_formats: raise gr.Error("Please select at least one output format.")
    
    chat_deployment = AZURE_DEPLOYMENT_NAMES.get(chat_model)
    transcribe_deployment = AZURE_DEPLOYMENT_NAMES.get(transcribe_model)
    if not chat_deployment: raise gr.Error(f"Deployment for model '{chat_model}' not found in configuration.")

    shutil.rmtree(OUTPUT_DIR, ignore_errors=True); shutil.rmtree(PREVIEW_DIR, ignore_errors=True)
    OUTPUT_DIR.mkdir(); PREVIEW_DIR.mkdir()

    # --- Step 1: Omni-Model Processing ---
    md_content = ""
    # Process files first
    if omni_files:
        # Check for multiple file types
        file_paths = [Path(f.name) for f in omni_files]
        extensions = {p.suffix.lower() for p in file_paths}

        if '.md' in extensions:
            md_content = "\n\n".join([p.read_text(encoding='utf-8') for p in file_paths if p.suffix.lower() == '.md'])
        elif '.pdf' in extensions:
             # For simplicity, we process only the first PDF if multiple are uploaded for OCR
            pdf_file = next((f for f in omni_files if Path(f.name).suffix.lower() == '.pdf'), None)
            ocr_prompt = omni_prompt if omni_prompt else "Extract all text from the following document pages."
            md_content = process_pdf_input(pdf_file, ocr_prompt, chat_deployment, progress)
        elif '.png' in extensions or '.jpg' in extensions or '.jpeg' in extensions:
            image_file = next((f for f in omni_files if Path(f.name).suffix.lower() in ['.png', '.jpg', '.jpeg']), None)
            md_content = process_image_input(image_file, omni_prompt, chat_deployment)
        elif '.wav' in extensions or '.mp3' in extensions or '.m4a' in extensions:
            if not transcribe_deployment: raise gr.Error(f"Deployment for model '{transcribe_model}' not found.")
            audio_file = next((f for f in omni_files if Path(f.name).suffix.lower() in ['.wav', '.mp3', '.m4a']), None)
            md_content = process_audio_input(audio_file, omni_prompt, chat_deployment, transcribe_deployment)
    # If no files, process text prompt
    elif omni_prompt:
        md_content = process_text_input(omni_prompt, chat_deployment)
    
    if not md_content: raise gr.Error("Failed to generate source content from the provided input.")
    
    # --- Step 2: Generate Selected Document Formats ---
    generated_files = []
    for format_choice in progress.tqdm(output_formats, desc="Generating Formats"):
        time_str = datetime.datetime.now().strftime('%m-%d-%a_%I%M%p').upper()
        if format_choice == "PDF":
            for layout_name in layouts:
                for font_name in fonts:
                    pagesize = LAYOUTS[layout_name]["size"]
                    final_pagesize = (pagesize[0] * page_w_mult, pagesize[1] * page_h_mult)
                    pdf_buffer = create_pdf(md_content, font_name, EMOJI_FONT_NAME, final_pagesize, num_columns)
                    filename = f"Document_{time_str}_{layout_name.replace(' ','-')}_{font_name}.pdf"
                    output_path = OUTPUT_DIR / filename
                    with open(output_path, "wb") as f: f.write(pdf_buffer.getvalue())
                    generated_files.append(output_path)
        elif format_choice == "DOCX":
            docx_doc = create_docx(md_content)
            filename = f"Document_{time_str}.docx"
            output_path = OUTPUT_DIR / filename
            docx_doc.save(output_path); generated_files.append(output_path)
        elif format_choice == "XLSX":
            xlsx_book = create_xlsx(md_content)
            filename = f"Outline_{time_str}.xlsx"
            output_path = OUTPUT_DIR / filename
            xlsx_book.save(output_path); generated_files.append(output_path)
            
    gallery_previews = [create_pdf_preview(p) for p in generated_files if p.suffix == '.pdf']
    final_gallery = [g for g in gallery_previews if g]
    
    return md_content, final_gallery, [str(p) for p in generated_files]

# --- 🎨 Gradio UI Definition ---
AVAILABLE_FONTS, EMOJI_FONT_NAME = register_local_fonts()

with gr.Blocks(theme=gr.themes.Soft(), title="Omni-Model Document Generator") as demo:
    gr.Markdown("# 🧠 Omni-Model Document Generator (PDF, DOCX, XLSX)")
    gr.Markdown("Provide a prompt, or upload a Markdown, PDF, Image, or Audio file. The AI will process it, and you can generate documents from the result.")

    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### ⚙️ Omni-Model Input")
            
            chat_models = ["gpt-4o", "gpt-4.1", "gpt-4.1-mini", "gpt-4o-mini", "o1-mini", "o3-mini", "o4-mini"]
            transcribe_models = ["gpt-4o-transcribe", "gpt-4o-mini-transcribe"]
            
            selected_chat_model = gr.Dropdown(choices=chat_models, label="Select Chat/Vision/Reasoning Model", value=chat_models[0])
            selected_transcribe_model = gr.Dropdown(choices=transcribe_models, label="Select Transcription Model (for audio)", value=transcribe_models[0])

            omni_prompt = gr.Textbox(label="Prompt", lines=3, placeholder="Ask a question, or provide instructions for a file...")
            omni_files = gr.File(label="Upload File(s) (Optional)", file_count="multiple", file_types=["image", ".wav", ".mp3", ".md", ".pdf"])
            
            gr.Markdown("### 📄 Output Settings")
            output_formats = gr.CheckboxGroup(choices=["PDF", "DOCX", "XLSX"], label="Select Output Formats", value=["PDF"])
            
            with gr.Accordion("PDF Customization", open=True):
                num_columns_slider = gr.Slider(label="Text Columns", minimum=1, maximum=4, step=1, value=1)
                page_w_mult_slider = gr.Slider(label="Page Width Multiplier", minimum=1, maximum=5, step=1, value=1)
                page_h_mult_slider = gr.Slider(label="Page Height Multiplier", minimum=1, maximum=2, step=1, value=1)
                selected_layouts = gr.CheckboxGroup(choices=list(LAYOUTS.keys()), label="Base Page Layout", value=["A4 Portrait"])
                selected_fonts = gr.CheckboxGroup(choices=AVAILABLE_FONTS, label="Text Font", value=[AVAILABLE_FONTS[0]] if AVAILABLE_FONTS else [])
            
            generate_btn = gr.Button("🚀 Generate Documents", variant="primary")
        
        with gr.Column(scale=2):
            gr.Markdown("### 🤖 AI Response (Source for Documents)")
            ai_response_output = gr.Markdown(label="AI Generated Content")
            gr.Markdown("### 🖼️ Final Documents")
            gallery_output = gr.Gallery(label="PDF Previews", show_label=False, elem_id="gallery", columns=3, height="auto", object_fit="contain")
            downloadable_files_output = gr.Files(label="Download Generated Files")
            
    generate_btn.click(fn=generate_outputs_api, 
                       inputs=[omni_files, omni_prompt, selected_chat_model, selected_transcribe_model, output_formats, selected_layouts, selected_fonts, num_columns_slider, page_w_mult_slider, page_h_mult_slider], 
                       outputs=[ai_response_output, gallery_output, downloadable_files_output])

if __name__ == "__main__":
    demo.launch()
